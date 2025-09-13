import logging
import os
import random
from datetime import datetime, timedelta
from io import BytesIO

import httpx
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.config import Config
from app.services.file_service import iter_doc_elements, process_docx
from app.services.leasing_service import REPORTS_DIR

logger = logging.getLogger(__name__)

DELTA_API_URL = "https://service.deltasecurity.ru/api2"
DELTA_SEARCH_URL = "https://sb.deltasecurity.ru/search/contractors"
DELTA_AUTH_URL = "https://sb.deltasecurity.ru/profile/auth-ajax"
DELTA_REPORT_URL = "https://sb.deltasecurity.ru/report/get-summary-report"
RAW_REPORTS_DIR = os.path.join(REPORTS_DIR, "raw")


async def get_delta_session() -> str:
    """
    Выполняет вход на sb.deltasecurity.ru для получения актуального PHPSESSID.
    Кэширует сессию, чтобы не логиниться при каждом запросе.
    """
    global session_cache

    logger.info("Авторизация в 'Дельта Безопасность' для получения новой сессии...")

    if not Config.DELTA_SB_LOGIN or not Config.DELTA_SB_PASSWORD:
        raise HTTPException(
            status_code=500,
            detail="Учетные данные DELTA_SB_LOGIN и DELTA_SB_PASSWORD не настроены в .env",
        )

    payload = {
        "user[login]": Config.DELTA_SB_LOGIN,
        "user[password]": Config.DELTA_SB_PASSWORD,
        "user[remember]": "1",
    }
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/108.0.0.0 Safari/537.36",
    }

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                DELTA_AUTH_URL,
                data=payload,
                headers=headers,
                timeout=20.0,
                follow_redirects=True,
            )
            response.raise_for_status()

            session_id = response.cookies.get("PHPSESSID")
            if not session_id:
                logger.error(
                    f"Не удалось получить PHPSESSID. Ответ сервера: {response.text}"
                )
                raise HTTPException(
                    status_code=500,
                    detail="Ошибка авторизации в Дельта: не удалось получить PHPSESSID. Проверьте учетные данные.",
                )
            print("session_id", session_id)
            cookie_str = f"PHPSESSID={session_id}"
            session_cache["cookie"] = cookie_str
            session_cache["expires_at"] = datetime.utcnow() + timedelta(days=1)

            logger.info("Успешная авторизация, PHPSESSID получен и кэширован.")
            return cookie_str

        except httpx.HTTPStatusError as e:
            logger.error(f"Ошибка HTTP при авторизации в Дельта: {e.response.text}")
            raise HTTPException(
                status_code=e.response.status_code,
                detail="Ошибка сервиса Дельта при авторизации.",
            )
        except Exception as e:
            logger.error(f"Непредвиденная ошибка при авторизации в Дельта: {e}")
            raise HTTPException(
                status_code=500, detail="Внутренняя ошибка при авторизации в Дельта."
            )


async def find_delta_company_id(inn: str) -> int:
    """
    Находит внутренний ID компании в 'Дельта Безопасность' по ИНН.
    """
    session_cookie = await get_delta_session()

    url = f"{DELTA_SEARCH_URL}?parts=snippet,highlight"
    payload = {
        "contractorType": "company",
        "keyword": inn,
        "page": 1,
        "resultsPerPage": 1,
        "filters": {
            "region_id": [],
            "city_id": [],
            "find_in": [],
            "activity_code": [],
            "only_main_activity": False,
            "status_group": [],
            "kladr_code": "",
            "opf_group": [],
        },
    }
    headers = {
        "Accept": "*/*",
        "Content-Type": "application/json; charset=utf-8",
        "Cookie": session_cookie,
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/108.0.0.0 Safari/537.36",
    }

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                url, headers=headers, json=payload, timeout=20.0
            )
            response.raise_for_status()
            data = response.json()

            if data.get("items") and len(data["items"]) > 0:
                company_id = data["items"][0].get("id")
                if company_id:
                    logger.info(f"Найден ID {company_id} для ИНН {inn} в Дельта.")
                    return company_id
                else:
                    raise HTTPException(
                        status_code=404,
                        detail=f"ID не найден в ответе от Дельта для ИНН {inn}",
                    )
            else:
                raise HTTPException(
                    status_code=404,
                    detail=f"Контрагент с ИНН {inn} не найден в Дельта.",
                )
        except httpx.HTTPStatusError as e:
            logger.error(
                f"Ошибка HTTP при поиске ID в Дельта для ИНН {inn}: {e.response.text}"
            )
            raise HTTPException(
                status_code=e.response.status_code,
                detail=f"Ошибка сервиса Дельта при поиске ID: {e.response.text}",
            )
        except Exception as e:
            logger.error(
                f"Непредвиденная ошибка при поиске ID в Дельта для ИНН {inn}: {e}"
            )
            raise HTTPException(
                status_code=500,
                detail=f"Внутренняя ошибка при поиске ID в Дельта: {str(e)}",
            )


async def get_processed_delta_report(inn: str, db: Session) -> tuple[str, str]:
    """
    Скачивает, очищает, векторизует и сохраняет отчет из Дельты.
    Возвращает путь к файлу и его текстовое содержимое.
    """
    if not Config.DELTA_SB_LOGIN or not Config.DELTA_SB_PASSWORD:
        raise HTTPException(status_code=500, detail="Учетные данные для Дельты не настроены.")

    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
        try:
            # --- Шаг 1: Авторизация ---
            logger.info("Авторизация в 'Дельта Безопасность'...")
            auth_payload = {
                "user[login]": Config.DELTA_SB_LOGIN,
                "user[password]": Config.DELTA_SB_PASSWORD,
                "user[remember]": "1"
            }
            auth_response = await client.post(DELTA_AUTH_URL, data=auth_payload)
            auth_response.raise_for_status()
            if "PHPSESSID" not in auth_response.cookies:
                raise HTTPException(status_code=401, detail="Авторизация в Дельте не удалась. PHPSESSID не получен.")
            logger.info("Авторизация успешна.")

            # --- Шаг 2: Поиск ID компании ---
            logger.info(f"Поиск ID для ИНН {inn}...")
            search_url = f"{DELTA_SEARCH_URL}?parts=snippet,highlight"
            search_payload = {"contractorType": "company", "keyword": inn, "page": 1, "resultsPerPage": 1,
                              "filters": {}}
            search_response = await client.post(search_url, json=search_payload)
            search_response.raise_for_status()
            search_data = search_response.json()
            items = search_data.get("items")
            if not items:
                raise HTTPException(status_code=404, detail=f"Контрагент с ИНН {inn} не найден в Дельта.")
            company_id = items[0].get("id")
            if not company_id:
                raise HTTPException(status_code=404, detail="Не удалось извлечь ID компании.")
            logger.info(f"Найден ID: {company_id}")

            # --- Шаг 3: Запрос отчета ---
            logger.info(f"Запрос отчета для компании ID {company_id}...")
            request_id = random.randint(10000000, 99999999)
            sources = "due-diligence,analitic,167,170,80,financial-info,227"
            report_url = f"{DELTA_REPORT_URL}/{company_id}/company/7?request_id={request_id}&sources={sources}"
            report_response = await client.get(report_url)
            report_response.raise_for_status()

            # --- Шаг 4: Обработка отчета ---
            logger.info("Отчет получен, начинаю обработку...")
            original_doc = Document(BytesIO(report_response.content))
            new_doc = Document()
            report_texts = []

            stop_processing = False
            for element in iter_doc_elements(original_doc):
                text = ""
                if isinstance(element, Paragraph):
                    text = element.text.strip()
                elif isinstance(element, Table) and element.rows and element.columns:
                    text = element.cell(0, 0).text.strip()

                if "Показатель (в руб.)" in text:
                    stop_processing = True
                if stop_processing:
                    continue

                # Копируем только нужные секции
                # (здесь можно добавить вашу логику по секциям, если она нужна)
                if isinstance(element, Paragraph):
                    new_doc.add_paragraph(element.text)
                    report_texts.append(element.text)
                elif isinstance(element, Table):
                    new_table = new_doc.add_table(rows=0, cols=len(element.columns))
                    new_table.style = element.style
                    for row in element.rows:
                        new_row = new_table.add_row()
                        row_text_parts = []
                        for i, cell in enumerate(row.cells):
                            if i < len(new_row.cells):
                                new_row.cells[i].text = cell.text
                                row_text_parts.append(cell.text.strip())
                        report_texts.append(" | ".join(row_text_parts))
                    new_doc.add_paragraph()

            # --- Шаг 5: Сохранение и векторизация ---
            clean_content_io = BytesIO()
            new_doc.save(clean_content_io)
            clean_content = clean_content_io.getvalue()

            clean_filename = f"delta_report_clean_{inn}.docx"
            file_path_in_db = "delta_reports"

            clean_upload_file = UploadFile(filename=clean_filename, file=BytesIO(clean_content))
            await process_docx(clean_upload_file, file_path_in_db, db)

            os.makedirs(REPORTS_DIR, exist_ok=True)
            final_path = os.path.join(REPORTS_DIR, clean_filename)
            with open(final_path, "wb") as f:
                f.write(clean_content)

            logger.info(f"Очищенный отчет для ИНН {inn} сохранен и векторизован.")

            # ✅ Возвращаем и путь, и текстовое содержимое
            return final_path, "\n".join(report_texts)

        except httpx.HTTPStatusError as e:
            error_text = e.response.text[:500]
            logger.error(f"Ошибка HTTP при работе с Дельтой: {e.response.status_code} - {error_text}")
            raise HTTPException(status_code=e.response.status_code, detail=f"Ошибка сервиса Дельта: {error_text}")
        except Exception as e:
            logger.error(f"Непредвиденная ошибка при работе с Дельтой для ИНН {inn}: {e}")
            raise HTTPException(status_code=500, detail=str(e))
