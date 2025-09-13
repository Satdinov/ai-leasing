import logging
import os
from io import BytesIO

import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi import HTTPException, UploadFile, status
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from sqlalchemy.orm import Session

from app.config import Config
from app.database import engine
from app.models import FileMetadata
from app.utils import normalize_table_name

logger = logging.getLogger(__name__)

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=Config.CHUNK_SIZE, chunk_overlap=Config.CHUNK_OVERLAP
)

vectorstore = Chroma(
    collection_name="docx_data",
    persist_directory=Config.PERSIST_DIRECTORY,
    embedding_function=OpenAIEmbeddings(
        model="text-embedding-3-small",
        api_key=os.getenv("AITUNNEL_API_KEY"),
        base_url="https://api.aitunnel.ru/v1/",
    ),
)


async def process_xlsx(file: UploadFile, file_path: str = "", db: Session = None):
    temp_path = f"temp_{file.filename}"
    file_name = file.filename
    try:
        with open(temp_path, "wb") as f:
            content = await file.read()
            if len(content) > Config.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail="File size exceeds maximum allowed size",
                )
            f.write(content)
        df = pd.read_excel(temp_path, engine="openpyxl")
        table_name = normalize_table_name(file.filename)
        try:
            file_name = file.filename.encode("cp437").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            file_name = file.filename.encode("utf-8").decode("utf-8")
        df.to_sql(name=table_name, con=engine, if_exists="replace", index=False)
        file_metadata = FileMetadata(
            file_name=file_name,
            file_path=file_path,
            file_type="xlsx",
            table_name=table_name,
            file_size=len(content),
        )
        db.add(file_metadata)
        db.commit()
        return {
            "status": "success",
            "filename": file_name,
            "table_name": table_name,
            "path": file_path,
        }
    except Exception as e:
        logger.error(f"Error processing XLSX {file_name}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing {file_name}: {str(e)}",
        )
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def iter_doc_elements(document):
    """
    Генератор, который возвращает параграфы и таблицы в том порядке, в котором они появляются в документе.
    """
    for element in document.element.body:
        if element.tag.endswith("p"):
            yield Paragraph(element, document)
        elif element.tag.endswith("tbl"):
            yield Table(element, document)


async def process_docx(file: UploadFile, file_path: str = "", db: Session = None):
    temp_path = f"temp_{file.filename}"
    file_name = file.filename
    try:
        content = await file.read()
        if len(content) > Config.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail="File size exceeds maximum allowed size",
            )

        doc = Document(BytesIO(content))
        full_text = []

        for element in iter_doc_elements(doc):
            if isinstance(element, Paragraph):
                full_text.append(element.text)
            elif isinstance(element, Table):
                md_table_rows = []
                is_header = True
                for row in element.rows:
                    row_text = [
                        cell.text.strip().replace("\n", " ") for cell in row.cells
                    ]
                    md_table_rows.append("| " + " | ".join(row_text) + " |")

                    if is_header:
                        md_table_rows.append(
                            "| " + " | ".join(["---"] * len(row.cells)) + " |"
                        )
                        is_header = False

                if md_table_rows:
                    full_text.append("\n".join(md_table_rows))

        text_content = "\n".join(full_text)

        try:
            file_name = file.filename.encode("cp437").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            file_name = file.filename.encode("utf-8").decode("utf-8")

        split_texts = text_splitter.split_text(text_content)
        if split_texts:
            vectorstore.add_texts(
                texts=split_texts,
                metadatas=[
                    {"source": file_path, "filename": file_name, "file_type": "docx"}
                    for _ in split_texts
                ],
            )

        file_metadata = FileMetadata(
            file_name=file_name,
            file_path=file_path,
            file_type="docx",
            file_size=len(content),
        )
        db.add(file_metadata)
        db.commit()

        os.makedirs(Config.DOCUMENTS_DIR, exist_ok=True)
        text_path = os.path.join(
            Config.DOCUMENTS_DIR, f"{os.path.splitext(file_name)[0]}.txt"
        )
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(text_content)

        return {
            "status": "success",
            "filename": file_name,
            "path": file_path,
            "chunks": len(split_texts),
        }
    except Exception as e:
        logger.error(f"Error processing DOCX {file_name}: {e}")
        if file.filename.lower().endswith(".doc"):
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Не удалось обработать файл .doc '{file_name}'. "
                f"Формат может быть несовместим. Попробуйте пересохранить его в .docx. Ошибка: {str(e)}",
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing {file_name}: {str(e)}",
        )
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


async def process_delta_report(file: UploadFile, db: Session):
    """
    Обрабатывает отчет Дельты: извлекает ключевые разделы,
    создает новый 'чистый' документ и добавляет его в векторную базу.
    """
    try:
        content = await file.read()
        original_doc = Document(BytesIO(content))
        new_doc = Document()

        in_delta_analitika = False
        in_financial_info = False
        in_external_info = False

        stop_keyword = "Показатель (в руб.)"

        for element in iter_doc_elements(original_doc):
            text = ""
            if isinstance(element, Paragraph):
                text = element.text.strip()
            elif isinstance(element, Table):
                text = element.cell(0, 0).text.strip()

            if "ДельтаАналитика" in text:
                in_delta_analitika = True
            if "Финансовая информация" in text:
                in_financial_info = True
            if "Информация от внешних источников" in text:
                in_external_info = True
            if stop_keyword in text:
                in_financial_info = False

            if in_delta_analitika or in_financial_info or in_external_info:
                if isinstance(element, Paragraph):
                    new_doc.add_paragraph(element.text)
                elif isinstance(element, Table):
                    new_table = new_doc.add_table(rows=0, cols=len(element.columns))
                    for row in element.rows:
                        new_row = new_table.add_row()
                        for i, cell in enumerate(row.cells):
                            new_row.cells[i].text = cell.text
                    new_doc.add_paragraph()

        clean_content_io = BytesIO()
        new_doc.save(clean_content_io)
        clean_content = clean_content_io.getvalue()

        clean_filename = f"clean_{file.filename}"

        clean_upload_file = UploadFile(
            filename=clean_filename, file=BytesIO(clean_content)
        )

        file_path_in_db = "delta_reports"

        logger.info(f"Запуск векторизации для очищенного отчета: {clean_filename}")

        return await process_docx(clean_upload_file, file_path_in_db, db)

    except Exception as e:
        logger.error(
            f"Ошибка при специальной обработке отчета Дельты '{file.filename}': {e}"
        )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка при обработке отчета Дельты: {str(e)}",
        )
